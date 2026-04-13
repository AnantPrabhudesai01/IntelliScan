from docx import Document
from docx.shared import Pt, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn

def set_cell_border(cell, **kwargs):
    """
    Set cell border
    Usage:
    set_cell_border(
        cell,
        top={"sz": 12, "val": "single", "color": "#000000", "space": "0"},
        bottom={"sz": 12, "val": "single", "color": "#000000", "space": "0"},
        left={"sz": 12, "val": "single", "color": "#000000", "space": "0"},
        right={"sz": 12, "val": "single", "color": "#000000", "space": "0"},
    )
    """
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()

    for edge in ('top', 'start', 'bottom', 'end'):
        node_name = edge
        if edge == 'start': node_name = 'left'
        if edge == 'end': node_name = 'right'
        
        edge_data = kwargs.get(edge)
        if edge_data:
            tag = 'w:{}'.format(node_name)
            element = tcPr.find(qn(tag))
            if element is None:
                element = OxmlElement(tag)
                tcPr.append(element)

            for key, value in edge_data.items():
                element.set(qn('w:{}'.format(key)), str(value))

def create_report(month_num, month_name, duration, tasks):
    doc = Document()
    
    # Title
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run('Sardar Vallabhbhai Patel Institute of Technology, Vasad')
    run.bold = True
    run.font.size = Pt(14)
    
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run('MCA (Semester - 4th)')
    run.bold = True
    run.font.size = Pt(12)
    
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run('Subject Code & Name: MC04094011 Research Work(Phase-2)/Major Project')
    run.bold = True
    run.font.size = Pt(11)
    
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(f'Monthly Progress Report-{month_num}')
    run.bold = True
    run.font.size = Pt(14)
    run.font.underline = True
    
    # Project Information
    doc.add_paragraph()
    run = doc.add_paragraph().add_run('Project Information')
    run.bold = True
    run.font.size = Pt(12)
    
    info = [
        ('Enrollment No. & Name :', ' [Your Enrollment No], Anant Prabhudesai'),
        ('Company Name :', ' IntelliScan'),
        ('Project Title :', ' IntelliScan: AI-Powered CRM & Business Card Management'),
        ('Progress Report Duration :', f' {duration}')
    ]
    
    for label, val in info:
        p = doc.add_paragraph()
        run_l = p.add_run(label)
        run_l.bold = True
        p.add_run(val)
        
    p = doc.add_paragraph()
    p.add_run('Project External Guide Name :').bold = True
    p.add_run(' [External Guide Name]')
    
    p = doc.add_paragraph()
    p.add_run('Project Internal Guide Name :').bold = True
    p.add_run(' [Internal Guide Name]')
    
    # Work Progress Table
    doc.add_paragraph()
    run = doc.add_paragraph().add_run('Work Progress Status')
    run.bold = True
    run.font.size = Pt(12)
    
    table = doc.add_table(rows=1, cols=4)
    table.style = 'Table Grid'
    hdr_cells = table.rows[0].cells
    hdr_cells[0].text = 'Date'
    hdr_cells[1].text = 'Task'
    hdr_cells[2].text = 'Task Status (In progress/ Completed)'
    hdr_cells[3].text = 'Remark by External Guide'
    
    for cell in hdr_cells:
        for p in cell.paragraphs:
            for run in p.runs:
                run.bold = True
                run.font.size = Pt(10)

    for date, task, status in tasks:
        row_cells = table.add_row().cells
        row_cells[0].text = date
        row_cells[1].text = task
        row_cells[2].text = status
        row_cells[3].text = '' # Remark empty

    # Signatures
    doc.add_paragraph()
    doc.add_paragraph()
    doc.add_paragraph()
    
    p = doc.add_paragraph()
    p.add_run('Student’s Enrollment No, Student’s Name & Signature:').bold = True
    
    doc.add_paragraph()
    p = doc.add_paragraph()
    p.add_run('External Guide’s Name & Signature:').bold = True
    
    doc.add_paragraph()
    p = doc.add_paragraph()
    p.add_run('Company’s Stamp:').bold = True
    
    filename = f'Monthly_Progress_Report_{month_num}_{month_name}.docx'
    doc.save(filename)
    return filename

# Data Preparation
reports_data = [
    {
        'num': 1, 'name': 'Jan', 'duration': 'January 1st, 2026 to January 31st, 2026',
        'tasks': [
            ('01/01 - 07/01', 'Requirement gathering and detailed market gap analysis.', 'Completed'),
            ('08/01 - 15/01', 'Finalizing technical stack (Node.js, SQLite, React, Gemini AI).', 'Completed'),
            ('16/01 - 22/01', 'Designing comprehensive UML diagrams (Use Case, Class, Activity).', 'Completed'),
            ('23/01 - 31/01', 'Preparing System Requirement Specification (SRS) documentation.', 'Completed')
        ]
    },
    {
        'num': 2, 'name': 'Feb', 'duration': 'February 1st, 2026 to February 28th, 2026',
        'tasks': [
            ('01/02 - 10/02', 'Initializing Backend environment and SQLite database orchestration.', 'Completed'),
            ('11/02 - 20/02', 'Implementing JWT Authentication and Tier-based Access Control.', 'Completed'),
            ('21/02 - 28/02', 'Developing core API endpoints for Profile and User management.', 'Completed')
        ]
    },
    {
        'num': 3, 'name': 'Mar', 'duration': 'March 1st, 2026 to March 31st, 2026',
        'tasks': [
            ('01/03 - 10/03', 'Mass migration of 60+ legacy HTML pages into routable React components.', 'Completed'),
            ('11/03 - 20/03', 'Integrating Google Gemini Multimodal API for intelligent card scanning.', 'Completed'),
            ('21/03 - 31/03', 'Implementing CRM Pipeline logic and AI Copywriter drafts.', 'Completed')
        ]
    },
    {
        'num': 4, 'name': 'Apr', 'duration': 'April 1st, 2026 to April 4th, 2026',
        'tasks': [
            ('01/04 - 02/04', 'Stabilizing navigation routes and fixing 500 server errors.', 'Completed'),
            ('03/04 - 04/04', 'Implementing Meeting Tools Download and Profile Persistence logic.', 'Completed'),
            ('Ongoing', 'Finalizing AI Outreach Sequence workflows and permission gating.', 'In Progress')
        ]
    }
]

generated_files = []
for r in reports_data:
    fname = create_report(r['num'], r['name'], r['duration'], r['tasks'])
    generated_files.append(fname)

print("Generated files:")
for f in generated_files:
    print(f"- {f}")
